import sys
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QMainWindow, QInputDialog, QLineEdit, QMessageBox,QFileDialog
from PyQt5.uic import loadUi
from encryption import Encrypt
from decryption import Decrypt
import docx
import os
class Core(QMainWindow):

    def __init__(self):
        super(Core, self).__init__()
        loadUi('GUI.ui',self)
        self.encrypt_btn.clicked.connect(self.encrypt)
        self.decrypt_btn.clicked.connect(self.decrypt)
        self.browse_file.clicked.connect(self.openFile)
        self.encrypt_file.clicked.connect(self.encryptFile)
        self.decrypt_file.clicked.connect(self.decryptFile)
        
    def encrypt(self):
        #
        self.plaintext = self.text_edit.toPlainText()
        if(self.plaintext!=""):
            self.password,ok = QInputDialog.getText(self,"Password","Kindly Enter the encryption password: ",QLineEdit.Password,"")
            if(len(self.password)<8  or self.password.islower() or self.password.isupper()):
                QMessageBox.about(self, "Password Error", "Password Length must be 8 or more characters and contain upper and lower case characters")
            else:
                enc = Encrypt(self.plaintext,self.password,'text')
                encrypted_text = enc.get_encrypted_bin()
                if enc:
                    QMessageBox.about(self, "Encryption Status", "Encryption Completed successfully")
                    self.text_edit.clear()
                    
                    self.text_edit.append(encrypted_text)
        else:
            QMessageBox.about(self,"Error","Please input a text to encrypt")
    def decrypt(self):
        #
        self.ciphertext = self.text_edit.toPlainText()
        if(self.ciphertext!=""):
            self.password,ok = QInputDialog.getText(self,"Password","Kindly Enter the encryption password: ",QLineEdit.Password,"")
            dec = Decrypt(self.ciphertext,self.password)
            decrypted_text = dec.getDecryptedString()
            if dec:
                QMessageBox.about(self, "Decryption Status", "Decryption Completed successfully")
                self.text_edit.clear()
                
                self.text_edit.append(decrypted_text)
        else:
            QMessageBox.about(self,"Error","Please input a text to decrypt")
    def openFile(self):
        self.fileName = QFileDialog.getOpenFileName(self)
        self.file_path = self.fileName[0]
        if(self.file_path!=""):            
            self.filename_field.setText(self.file_path)
            self.filetype = self.file_path.rpartition('.')[-1]
    def encryptFile(self):        
        if(self.file_path!=""):            
            # self.filename_field.setText(self.file_path)
            self.password,ok = QInputDialog.getText(self,"Password","Kindly Enter the encryption password: ",QLineEdit.Password,"")
            #read the file
            if(len(self.password)<8  or self.password.islower() or self.password.isupper()):
                QMessageBox.about(self, "Password Error", "Password Length must be 8 or more characters and contain upper and lower case characters")
            else:
                if(self.filetype == 'txt'):
                    file_r = open(self.file_path,"r")
                    
                    file_data = file_r.read()
                    
                    print(file_data)
                    
                    
                    fe = Encrypt(str(file_data),self.password,"text")
                    enc_bin = fe.get_encrypted_bin()
                    
                    file_wr =open(self.file_path,"w")
                    file_wr.write(enc_bin)
                    file_wr.close()
                    if(file_wr):
                        QMessageBox.about(self, "Encryption Status", "File encryption Completed successfully")
                        
                        self.file_enc_field.append("File Location " +self.file_path)
                elif(self.filetype == 'doc' or self.filetype == 'docx'):
                    doc = docx.Document(self.file_path)

                    file_data = ""
                    
                    paragraphs = doc.paragraphs
                    
                    for paragraph in paragraphs:
                        file_data+=paragraph.text
                    #encrypt the read file data
                    fe = Encrypt(str(file_data),self.password,"text")
                    enc_bin = fe.get_encrypted_bin()
                    
                    #create a new file to save the encrypted info
                    new_doc =  docx.Document()
                    
                    new_doc.add_paragraph(enc_bin)
                    
                    new_doc.save(self.file_path)
                    if(new_doc):                    
                        QMessageBox.about(self, "Encryption Status", "File encryption Completed successfully")
                        self.file_enc_field.clear()
                        self.file_enc_field.append("File Path " +self.file_path) 
                    
                        print("successful encryption")                               
                else:
                    file_r = open(self.file_path,"rb")
                    file = file_r.read()
                    file = str(file[2:]).replace("b'","")
                    print(file[:20])
                    en = Encrypt(str(file),self.password,"file")

                    encrypted_text = en.get_encrypted_bin()

                    print(encrypted_text[:30])

                    fin = open(self.file_path+".enc", 'w')

                    # encrypted_text = bytearray(encrypted_text,"utf-8")
                    # writing decryption data in image
                    fin.write(encrypted_text)
                    fin.close()
                    # os.remove(self.file_path)
                                            
        else:
            QMessageBox.about(self, "Error", "You have not selected a file")

    def decryptFile(self):
        if(self.file_path!=""):            
            # self.filename_field.setText(self.file_path)
            self.password,ok = QInputDialog.getText(self,"Password","Kindly Enter the encryption password: ",QLineEdit.Password,"")
            #read the file
            if(self.filetype == 'txt'):
                
                file_r = open(self.file_path,"r")
                
                file_data = file_r.read()
                
                print(file_data[:20])
                
                
                fe = Decrypt(str(file_data),self.password)
                dec_bin = fe.getDecryptedString()
                
                file_wr =open(self.file_path,"w")
                file_wr.write(dec_bin)
                file_wr.close()
                if(file_wr):
                    QMessageBox.about(self, "Decryption Status", "File decryption Completed successfully")
                    self.file_enc_field.clear()
                    self.file_enc_field.append("File Path " +self.file_path)
            elif(self.filetype == 'doc' or self.filetype == 'docx'):
                    doc = docx.Document(self.file_path)

                    file_data = ""
                    
                    paragraphs = doc.paragraphs
                    
                    for paragraph in paragraphs:
                        file_data+=paragraph.text
                    #decrypt the read file data
                    fe = Decrypt(str(file_data),self.password)
                    dec_bin = fe.getDecryptedString()
                    
                    #create a new file to save the encrypted info
                    new_doc =  docx.Document()
                    
                    new_doc.add_paragraph(dec_bin)
                    
                    new_doc.save(self.file_path)
                    if(new_doc):                    
                        QMessageBox.about(self, "Decryption Status", "File decryption Completed successfully")
                        self.file_enc_field.clear()
                        self.file_enc_field.append("File Path " +self.file_path)
                        
                        print("successful decryption")
            else:
                file_r = open(self.file_path,"r")
                file = file_r.read()

                
                dec = Decrypt(file,self.password,"file")

                decrypted_text = dec.getDecryptedString()

                print(decrypted_text[:20])
                data = int(decrypted_text, 2).to_bytes((len(decrypted_text) + 7) // 8, byteorder='big')

                self.file_path=self.file_path.replace(".enc","")

                fin = open(self.file_path, 'wb')
                
                

                print(decrypted_text[:20])
                # writing decryption data in image
                fin.write(data)
                fin.close()

                # encrypted_text = en.get_encrypted_bin() 
        else:
            QMessageBox.about(self, "Error", "You have not selected a file")

    
    
    
          
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)

    window = Core()
    window.show()
    sys.exit(app.exec())